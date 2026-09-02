#!/usr/bin/env python3
"""Générateur .docx du pack Meta Ads (livrable client).

Même branding que `vsl-copywriter/assets/build_vsl_docx.py` : noir sur blanc,
Calibri, aucun effet, aucune couleur d'accent. Le document doit pouvoir être imprimé,
annoté sur un plateau de tournage, et relu par un juriste.

Dépendance : python-docx (`pip3 install python-docx`).

Usage :
    python3 build_meta_ads_docx.py meta-ads/v1.md Meta-Ads-Client.docx \
        --client "Acme Conseil" --date "12 mars 2026" \
        --subtitle "La Méthode Signature — 90 jours" \
        --brand "Mon Agence"    # optionnel : nom affiché dans le pied de page

Markdown accepté (sous-ensemble volontairement restreint, identique au builder VSL) :
    # H1 · ## H2 · ### H3 · #### H4
    paragraphes, **gras**, *italique*, `code` (rendu en gras)
    - listes à puces · 1. listes numérotées
    > citations (rendues en encadré gris clair)
    | tableaux | pipe |
    --- (filet de séparation)

Conventions du pack Meta Ads reconnues et mises en forme spécialement :
    #### 1B — 60 secondes · 156 mots     → titre de script, budget de mots en tête
    **Hook** / **Body** / **CTA**        → étiquettes de bloc, petites capitales grises
    (les lignes qui suivent une étiquette sont la voix off, corps 11 pt)
    - **Primary text — court (≤125 c)** : …  → ligne de copie, corps 9,5 pt
    - **Headline (≤40 c)** : …               → ligne de copie
    *ligne entièrement en italique*       → indication scénique, 9 pt gris
    > **`[PREUVE À PRODUIRE]`** …         → encadré placeholder
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    sys.exit("python-docx manquant : pip3 install python-docx")

BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)
BOX_FILL = "F2F2F2"

# étiquettes de bloc de script : ce qui suit est de la voix off
VO_LABELS = ("**Hook**", "**Body**", "**CTA**")


# ---------------------------------------------------------------- utilitaires

def setup(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(2.2)


def shade(paragraph, fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "BFBFBF")
    pbdr.append(bottom)
    pPr.append(pbdr)


def footer(doc: Document, text: str) -> None:
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text + "  ·  ")
        r.font.size = Pt(8)
        r.font.color.rgb = GREY
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)


INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline(paragraph, text: str, size: float, color=BLACK, base_italic=False) -> None:
    """Ajoute `text` à `paragraph` en respectant **gras**, *italique* et `code`."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.size = Pt(size)
            r.font.color.rgb = color
            r.italic = base_italic
        seg = m.group(0)
        r = paragraph.add_run(seg[2:-2] if seg.startswith("**") else seg[1:-1])
        r.font.size = Pt(size)
        r.font.color.rgb = color
        if seg.startswith("**") or seg.startswith("`"):
            r.bold = True
            r.italic = base_italic
        else:
            r.italic = True
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.italic = base_italic


def strip_md(text: str) -> str:
    return re.sub(r"[*`]", "", text)


# ------------------------------------------------------------------- blocs

def heading(doc, text, size, space_before, space_after, color=BLACK, upper=False):
    p = doc.add_paragraph()
    r = p.add_run(strip_md(text).upper() if upper else strip_md(text))
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    return p


def label(doc, text):
    """Étiquette de bloc de script : Hook / Body / CTA."""
    p = doc.add_paragraph()
    r = p.add_run(strip_md(text).upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = GREY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    return p


def vo(doc, text):
    """Voix off d'un script : le texte réellement prononcé."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(4)
    add_inline(p, text, 11)
    return p


def quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    add_inline(p, text, 9.5)
    shade(p, BOX_FILL)
    return p


def table(doc, rows, size=9):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j in range(cols):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            val = row[j] if j < len(row) else ""
            add_inline(p, val, size)
            if i == 0:
                for r in p.runs:
                    r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ------------------------------------------------------------------ parseur

def render(md: str, doc: Document) -> None:
    lines = md.split("\n")
    i = 0
    pending_table: list[list[str]] = []
    in_script = False          # True entre une étiquette Hook/Body/CTA et la fin du bloc

    def flush_table():
        nonlocal pending_table
        if pending_table:
            table(doc, pending_table)
            pending_table = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # tableaux
        if stripped.startswith("|") and stripped.endswith("|"):
            in_script = False
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                pending_table.append(cells)
            i += 1
            continue
        flush_table()

        if not stripped:
            i += 1
            continue

        # filet
        if re.fullmatch(r"-{3,}", stripped):
            in_script = False
            rule(doc)
            i += 1
            continue

        # titres
        if stripped.startswith("#### "):
            in_script = False
            heading(doc, stripped[5:], 11, 12, 3)
            i += 1
            continue
        if stripped.startswith("### "):
            in_script = False
            heading(doc, stripped[4:], 12.5, 16, 4)
            i += 1
            continue
        if stripped.startswith("## "):
            in_script = False
            heading(doc, stripped[3:], 14, 18, 6)
            i += 1
            continue
        if stripped.startswith("# "):
            in_script = False
            heading(doc, stripped[2:], 19, 0, 8)
            i += 1
            continue

        # citation / encadré (plan média, réserves, placeholders de preuve)
        if stripped.startswith(">"):
            in_script = False
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            quote(doc, " ".join(b for b in buf if b))
            continue

        # étiquette de bloc de script
        if stripped in VO_LABELS:
            label(doc, stripped.strip("*"))
            in_script = True
            i += 1
            continue

        # listes
        if re.match(r"^[-*] ", stripped):
            in_script = False
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            # les lignes de copie Meta sont plus denses : corps réduit
            size = 9.5 if "Primary text" in stripped or "Headline" in stripped \
                or "Description" in stripped else 10
            add_inline(p, stripped[2:], size)
            i += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            in_script = False
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, re.sub(r"^\d+\.\s*", "", stripped), 10)
            i += 1
            continue

        # indication scénique (ligne entièrement en italique)
        if re.fullmatch(r"\*[^*].*[^*]\*", stripped):
            in_script = False
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run(stripped[1:-1])
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = GREY
            i += 1
            continue

        # voix off, ou paragraphe courant
        if in_script:
            vo(doc, stripped)
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped, 10.5)
        i += 1

    flush_table()


# --------------------------------------------------------------------- main

def build(src: Path, out: Path, client: str, date: str, subtitle: str | None, brand: str = "") -> None:
    md = src.read_text(encoding="utf-8")
    doc = Document()
    setup(doc)

    # entête : on remplace le H1 du markdown par un entête maîtrisé
    lines = md.split("\n")
    title = "Pack Meta Ads"
    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        md = "\n".join(lines[1:])

    heading(doc, title, 20, 0, 2)
    meta = f"{client} · {date}" + (f" · {subtitle}" if subtitle else "")
    p = doc.add_paragraph()
    r = p.add_run(meta)
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(10)
    rule(doc)

    render(md, doc)
    footer(doc, (f"{brand} — " if brand else "") + f"Confidentiel · {client}")
    doc.save(str(out))
    print("écrit :", out)


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Markdown de pack Meta Ads → .docx (livrable client)")
    ap.add_argument("source", type=Path)
    ap.add_argument("output", type=Path)
    ap.add_argument("--client", default="Client")
    ap.add_argument("--date", default="")
    ap.add_argument("--subtitle", default=None)
    ap.add_argument("--brand", default="", help="nom de l'agence affiché dans le pied de page (optionnel)")
    a = ap.parse_args()
    if not a.source.exists():
        sys.exit(f"introuvable : {a.source}")
    build(a.source, a.output, a.client, a.date, a.subtitle, a.brand)
